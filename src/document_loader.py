"""문서 로딩 및 파싱 모듈"""

import os
from pathlib import Path
from typing import List, Dict, Optional
import logging

import docx2txt
from docx import Document
import openpyxl
from PyPDF2 import PdfReader
from langchain_core.documents import Document as LangchainDocument

from .utils import get_all_documents, extract_category_from_path, clean_text

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentLoader:
    """문서 로더 클래스"""
    
    def __init__(self, root_dir: str):
        """
        Args:
            root_dir: 문서가 있는 루트 디렉토리
        """
        self.root_dir = root_dir
        self.supported_parsers = {
            '.docx': self._parse_docx,
            '.doc': self._parse_doc,
            '.xlsx': self._parse_xlsx,
            '.xls': self._parse_xlsx,
            '.pdf': self._parse_pdf,
        }
    
    def load_documents(self) -> List[LangchainDocument]:
        """
        모든 문서를 로드하고 파싱합니다.
        
        Returns:
            LangchainDocument 리스트
        """
        # zip 파일 제외하고 모든 문서 찾기
        file_paths = get_all_documents(self.root_dir, exclude_extensions=['.zip'])
        
        logger.info(f"총 {len(file_paths)}개의 문서를 발견했습니다.")
        
        documents = []
        failed_files = []
        
        for file_path in file_paths:
            try:
                doc = self._load_single_document(file_path)
                if doc and doc.page_content.strip():
                    documents.append(doc)
                    logger.info(f"✓ 로드 성공: {file_path.name}")
                else:
                    logger.warning(f"⚠ 내용 없음: {file_path.name}")
            except Exception as e:
                logger.error(f"✗ 로드 실패: {file_path.name} - {str(e)}")
                failed_files.append((file_path.name, str(e)))
        
        logger.info(f"\n=== 로딩 완료 ===")
        logger.info(f"성공: {len(documents)}개")
        logger.info(f"실패: {len(failed_files)}개")
        
        if failed_files:
            logger.warning("\n실패한 파일 목록:")
            for filename, error in failed_files:
                logger.warning(f"  - {filename}: {error}")
        
        return documents
    
    def _load_single_document(self, file_path: Path) -> Optional[LangchainDocument]:
        """
        단일 문서를 로드합니다.
        
        Args:
            file_path: 파일 경로
        
        Returns:
            LangchainDocument 또는 None
        """
        ext = file_path.suffix.lower()
        parser = self.supported_parsers.get(ext)
        
        if not parser:
            logger.warning(f"지원하지 않는 파일 형식: {ext}")
            return None
        
        # 텍스트 추출
        text = parser(file_path)
        
        if not text or not text.strip():
            return None
        
        # 텍스트 정리
        cleaned_text = clean_text(text)
        
        # 메타데이터 생성
        category = extract_category_from_path(file_path, self.root_dir)
        metadata = {
            'source': str(file_path),
            'filename': file_path.name,
            'category': category,
            'file_type': ext,
        }
        
        return LangchainDocument(
            page_content=cleaned_text,
            metadata=metadata
        )
    
    def _parse_docx(self, file_path: Path) -> str:
        """DOCX 파일 파싱"""
        try:
            # 방법 1: python-docx 사용
            doc = Document(str(file_path))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 표도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n'.join(text_parts)
        except:
            # 방법 2: docx2txt 사용 (fallback)
            try:
                return docx2txt.process(str(file_path))
            except Exception as e:
                logger.error(f"DOCX 파싱 실패: {file_path.name} - {str(e)}")
                return ""
    
    def _parse_doc(self, file_path: Path) -> str:
        """DOC 파일 파싱 (레거시 형식)"""
        try:
            # Windows에서는 win32com 사용
            import win32com.client
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            try:
                doc = word.Documents.Open(str(file_path.absolute()))
                text = doc.Content.Text
                doc.Close(False)
                return text
            finally:
                word.Quit()
        except ImportError:
            logger.warning(f"win32com을 사용할 수 없습니다. .doc 파일을 건너뜁니다: {file_path.name}")
            return ""
        except Exception as e:
            logger.error(f"DOC 파싱 실패: {file_path.name} - {str(e)}")
            return ""
    
    def _parse_xlsx(self, file_path: Path) -> str:
        """XLSX/XLS 파일 파싱"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"\n=== {sheet_name} ===\n")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"XLSX 파싱 실패: {file_path.name} - {str(e)}")
            return ""
    
    def _parse_pdf(self, file_path: Path) -> str:
        """PDF 파일 파싱"""
        try:
            reader = PdfReader(str(file_path))
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"\n--- Page {page_num + 1} ---\n")
                    text_parts.append(text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"PDF 파싱 실패: {file_path.name} - {str(e)}")
            return ""


def test_loader():
    """문서 로더 테스트 함수"""
    loader = DocumentLoader("./reference")
    documents = loader.load_documents()
    
    print(f"\n총 {len(documents)}개의 문서가 로드되었습니다.")
    
    if documents:
        print("\n첫 번째 문서 샘플:")
        print(f"파일명: {documents[0].metadata['filename']}")
        print(f"카테고리: {documents[0].metadata['category']}")
        print(f"내용 길이: {len(documents[0].page_content)} 글자")
        print(f"내용 미리보기:\n{documents[0].page_content[:500]}...")


if __name__ == "__main__":
    test_loader()

